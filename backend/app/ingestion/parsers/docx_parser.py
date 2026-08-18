"""DOCX text extraction via python-docx.

Heading paragraphs (Word styles "Heading 1".."Heading 6"/"Title") are
re-emitted with a Markdown-style "#" prefix so downstream chunking can
detect document structure without needing to touch python-docx objects.
"""

import io
import re

import docx

from app.ingestion.parsers.base import BaseDocumentParser, ParsedDocument

_HEADING_STYLE_RE = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)


def _markdown_prefix(style_name: str) -> str:
    if style_name.lower() == "title":
        return "# "
    match = _HEADING_STYLE_RE.match(style_name or "")
    if match:
        level = max(1, min(6, int(match.group(1))))
        return "#" * level + " "
    return ""


class DocxParser(BaseDocumentParser):
    def supports(self, filename: str) -> bool:
        return filename.lower().endswith(".docx")

    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        document = docx.Document(io.BytesIO(file_bytes))
        lines = []
        for p in document.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            prefix = _markdown_prefix(p.style.name if p.style else "")
            lines.append(f"{prefix}{text}" if prefix else text)
        full_text = "\n".join(lines)
        return ParsedDocument(
            text=full_text,
            pages=[full_text],
            metadata={"source_filename": filename, "page_count": 1, "format": "docx"},
        )
